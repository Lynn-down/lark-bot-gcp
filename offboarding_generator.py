"""
离职文档生成模块
- generate_resignation_certificate(fields) -> 离职证明路径
- generate_termination_agreement(fields)   -> 离职协议路径
- build_offboarding_email(fields)          -> (subject, body)
"""
import os
import re
import logging
from datetime import datetime
from docx import Document

logger = logging.getLogger(__name__)

_HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.environ.get("CONTRACT_TEMPLATES_DIR", os.path.join(_HERE, "templates"))
OUTPUT_DIR    = os.environ.get("CONTRACT_OUTPUT_DIR", "/tmp")

TPL_CERT      = os.path.join(TEMPLATES_DIR, "离职证明-模板.docx")
TPL_AGREEMENT = os.path.join(TEMPLATES_DIR, "协商一致解除劳动合同协议书--模板.docx")


# ── 工具函数 ──────────────────────────────────────────────────────────────────

def _fmt_cn(ds: str) -> str:
    """'2026-04-01' → '2026年4月1日'"""
    if not ds:
        return ""
    try:
        d = datetime.strptime(str(ds)[:10], "%Y-%m-%d")
        return f"{d.year}年{d.month}月{d.day}日"
    except Exception:
        return str(ds)


def _set_para_text(para, new_text: str):
    """把段落所有 run 内容合并到第一个 run，保留原来的字符格式外壳。"""
    runs = para.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def _replace_in_doc(doc, replacements: dict):
    """对整个文档（正文段落 + 表格单元格）做批量字符串替换。"""
    def _apply(para):
        full = para.text
        new = full
        for old, new_val in replacements.items():
            new = new.replace(old, new_val)
        if new != full:
            _set_para_text(para, new)

    for para in doc.paragraphs:
        _apply(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply(para)


# ── 离职证明 ──────────────────────────────────────────────────────────────────

def generate_resignation_certificate(fields: dict, output_name: str = None) -> str:
    """生成《离职证明》，返回文件路径。"""
    if not os.path.exists(TPL_CERT):
        raise FileNotFoundError(f"离职证明模板不存在: {TPL_CERT}")

    doc   = Document(TPL_CERT)
    today = datetime.now()
    name      = fields.get("name", "（员工姓名）")
    gender    = fields.get("gender", "（男/女）")
    id_number = fields.get("id_number", "（员工身份证号码）")
    start_cn  = _fmt_cn(fields.get("start_date", "")) or "（入职日期）"
    leave_cn  = _fmt_cn(fields.get("leave_date", "")) or "（离职日期）"
    dept      = fields.get("department", "（部门）")
    job_title = fields.get("job_title", "（职务）")
    reason    = fields.get("reason", "经协商一致解除合同")
    sign_cn   = f"{today.year}年{today.month}月{today.day}日"

    _replace_in_doc(doc, {
        "（员工姓名）":                    name,
        "（男/女）":                       gender,
        "（员工身份证号码）":              id_number,
        "（入职日期：YYYY年MM月DD日）":   start_cn,
        "（离职日期：YYYY年MM月DD日）":   leave_cn,
        "（部门名称，如：市场部）":        dept,
        "（具体职务，如：高级经理）":      job_title,
        "（请选择或填写离职原因）":        reason,
        "YYYY年MM月DD日":                 sign_cn,
    })

    fname    = output_name or name
    out_path = os.path.join(OUTPUT_DIR, f"{fname}-离职证明.docx")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(out_path)
    logger.info(f"离职证明已生成: {out_path}")
    return out_path


# ── 离职协议 ──────────────────────────────────────────────────────────────────

def generate_termination_agreement(fields: dict, output_name: str = None) -> str:
    """生成《协商一致解除劳动合同协议书》，返回文件路径。"""
    if not os.path.exists(TPL_AGREEMENT):
        raise FileNotFoundError(f"离职协议模板不存在: {TPL_AGREEMENT}")

    doc = Document(TPL_AGREEMENT)

    name      = fields.get("name", "")
    id_number = fields.get("id_number", "")
    phone     = fields.get("phone", "")
    compensation     = str(fields.get("compensation", ""))
    bank_account_name = fields.get("bank_account_name", name)  # 户名默认为本人姓名
    bank_name        = fields.get("bank_name", "")
    bank_account     = fields.get("bank_account", "")
    bank_branch      = fields.get("bank_branch", "")

    def _parse_dt(ds):
        try:
            return datetime.strptime(str(ds)[:10], "%Y-%m-%d")
        except Exception:
            return None

    start_dt = _parse_dt(fields.get("start_date", ""))
    leave_dt = _parse_dt(fields.get("leave_date", ""))

    for para in doc.paragraphs:
        t = para.text

        # 乙方信息块（多行段落）
        if "乙方（员工）：" in t:
            lines = t.split("\n")
            new_lines = []
            for line in lines:
                s = line.strip()
                if "乙方（员工）：" in s:
                    new_lines.append(f"乙方（员工）：{name}")
                elif s.startswith("身份证号："):
                    new_lines.append(f"身份证号：{id_number}")
                elif s.startswith("联系电话："):
                    new_lines.append(f"联系电话：{phone}")
                elif "入职日期：" in s and start_dt:
                    new_lines.append(
                        f"入职日期：{start_dt.year}年{start_dt.month}月{start_dt.day}日"
                    )
                elif "拟解除日期：" in s and leave_dt:
                    new_lines.append(
                        f"拟解除日期：{leave_dt.year}年{leave_dt.month}月{leave_dt.day}日"
                    )
                else:
                    new_lines.append(line)
            _set_para_text(para, "\n".join(new_lines))

        # 第 1.1 条：解除日填写
        elif "甲乙双方确认于" in t and leave_dt:
            new_t = re.sub(
                r"确认于\s+年\s+月\s+日",
                f"确认于{leave_dt.year}年{leave_dt.month}月{leave_dt.day}日",
                t,
            )
            if new_t != t:
                _set_para_text(para, new_t)

        # 甲方签署日期 = 拟解除日期（用户要求）
        elif "__________年____月____日" in t and leave_dt:
            new_t = t.replace(
                "__________年____月____日",
                f"{leave_dt.year}年{leave_dt.month}月{leave_dt.day}日",
            )
            _set_para_text(para, new_t)

    # 赔偿金、银行信息、支付期限：整文档替换
    comp_display = "无" if compensation in ("无", "0", "") else f"{compensation}元"
    payment_period = fields.get("payment_period", "1个月内")
    _replace_in_doc(doc, {
        "（经济补偿金金额）":   comp_display,
        "（支付期限）":         payment_period,
        "（户名）":             bank_account_name,
        "（开户行）":           bank_name or bank_branch,
        "（账号）":             bank_account,
    })

    fname    = output_name or name
    out_path = os.path.join(OUTPUT_DIR, f"{fname}-离职协议.docx")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(out_path)
    logger.info(f"离职协议已生成: {out_path}")
    return out_path


# ── 邮件内容构建 ──────────────────────────────────────────────────────────────

def build_offboarding_email(fields: dict) -> tuple:
    """
    返回 (subject, body_text)
    说明行按用户要求修改：岗位匹配度问题。
    """
    name      = fields.get("name", "您")
    leave_cn  = _fmt_cn(fields.get("leave_date", "")) or "约定日期"

    subject = f"关于离职手续办理的通知"
    body = (
        f"{name}：\n"
        f"你好！\n\n"
        f"经过我们的内部沟通，决定目前终止与您的合作，主要原因是岗位匹配度问题，"
        f"与您的个人能力无关。感谢您这几天在公司的付出，祝您一切顺利！\n\n"
        f"请在离职前注意以下事项：\n"
        f"  工作交接：请将手中负责的文档、账号及未竟事项整理并交接；\n"
        f"  资产归还：最后工作日（{leave_cn}）下班前，您的 Lark 应用等权限将关闭；\n"
        f"  流程办理：配合 HR 完成离职手续。\n\n"
        f"北京极群｜HR"
    )
    return subject, body
