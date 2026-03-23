#!/usr/bin/env python3
"""
严格按模板生成合同 - 只替换下划线部分
"""
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
import os

# 劳动合同模板 - 严格按照原模板，只替换___部分
LABOR_CONTRACT_TEMPLATE = """劳动合同（通用）

甲    方：北京极群科技有限公司
乙    方：___员工姓名___
签订日期：  ___签订年份___  年  ___签订月份___  月  ___签订日期___  日

劳动合同

本劳动合同（下称"本合同"）由以下双方于  ___签订年份___  年  ___签订月份___ 月  ___签订日期___  日签订：

甲方：
注册地址：北京市海淀区中关村东路8号东升大厦AB座四层4161号
法定代表人：陈春宇

乙方：
身份证号码：___身份证号___
户籍地址：___户籍地址___
联系地址：___联系地址___
联系电话：___手机号___

鉴于：
甲方已如实告知并经乙方确认其工作内容、工作条件、工作地点、职业危害、安全生产状况、劳动报酬以及乙方要求了解的其他情况；根据《中华人民共和国劳动法》《中华人民共和国劳动合同法》等法律法规政策规定,甲乙双方遵循合法、公平、平等自愿、协商一致、诚实信用的原则订立本合同，共同遵守所列条款。

合同期限
1.  本合同为（任选一种）：
（1）□无固定期限合同；
（2）☑有固定期限，合同期限  ___合同年限___  年，自  ___签订年份___  年 ___签订月份___ 月 ___签订日期___  日起至  ___结束年份___  年 ___结束月份___ 月 ___结束日期___  日止，试用期为 ___试用期月数___ 个月；
（3）□以完成一定工作任务为期限。

2.  本合同到期后，甲方未提出续签或乙方不同意续签的，本合同到期终止。

工作内容、工作地点
1.  乙方工作地点为     ___工作地点___      。

2.  签订合同时，乙方工作岗位为  ___岗位名称___   ，基本岗位职责详见附件1：基本岗位职责描述。

3.  乙方应当接受甲方对其进行的工作岗位所必需的培训。乙方应主动学习，积极参加甲方组织的培训，提高职业技能。

4.  乙方必须按照甲方确定的岗位职责，按时、按质、按量完成工作。

工作时间和休息休假
1.  甲方安排乙方执行标准工时工作制。每日工作时间8小时，每周工作时间40小时，每周至少休息一日。

2.  甲方安排乙方加班的，应依法优先安排补休，无法安排补休的依法支付加班工资。

3.  乙方依法享有国家规定的各种法定节假日和甲方的休假待遇。

四、劳动报酬
1.  甲方依法制定工资分配制度，并告知乙方。甲方至少每月以货币形式向乙方支付一次工资。甲方经与乙方协商，约定按照如下方式向乙方以货币形式发放工资，于每月15日前足额支付：
固定工资 ___税前工资___ 元/月（税前）。

2.  试用期内，乙方的工资为转正后的  100  %。

3.  双方按照国家和地方规定缴纳社会保险和住房公积金，甲方依法为乙方办理有关手续并履行代扣代缴义务。

五、岗位职责
乙方应履行的岗位职责为：
___岗位职责描述___

六、劳动纪律
1.  乙方应遵守国家的法律法规。

2.  乙方已知悉并详细阅读甲方的各项规章制度，并承诺严格遵守。甲方有权依法不时制定、修改内部规章制度，乙方亦有义务遵守。

3.  乙方应遵守的劳动纪律包括但不限于：
- 严格遵守考勤制度，不得迟到、早退、旷工；
- 服从甲方工作安排；
- 不得有弄虚作假和欺诈的行为；
- 保守甲方商业秘密等。

七、劳动保护、劳动条件和职业危害防护
甲方应按照国家有关规定为乙方提供必要的劳动条件和安全保护。

八、保密规定
在本合同有效期间及之后，乙方不得以任何方式侵害甲方和/或其关联方的商业秘密和/或保密信息，否则应当依法承担违约责任和其他一切法律责任。

九、合同的变更、解除和终止
按照国家相关法律法规执行。

十、争议处理
双方如发生劳动争议，应协商解决；协商不成的，可向劳动争议仲裁委员会申请仲裁。

十一、其他
1.  本合同一式两份，甲乙双方各执一份，具有同等法律效力。
2.  本合同自双方签字（或盖章）之日起生效。


甲方（盖章）：                    乙方（签字）：

日期：___签订年份___年___签订月份___月___签订日期___日        日期：___签订年份___年___签订月份___月___签订日期___日
"""


def generate_labor_contract(data: dict, output_path: str = None) -> str:
    """
    生成劳动合同 - 严格按照模板，只替换占位符
    """
    # 解析日期
    sign_date = data.get("签订日期", datetime.now().strftime("%Y-%m-%d"))
    sign_year, sign_month, sign_day = sign_date.split("-")
    
    # 计算结束日期
    contract_years = int(data.get("合同年限", "3"))
    end_date = datetime(int(sign_year), int(sign_month), int(sign_day)) + timedelta(days=365*contract_years)
    end_year, end_month, end_day = str(end_date.year), str(end_date.month).zfill(2), str(end_date.day).zfill(2)
    
    # 替换所有占位符
    content = LABOR_CONTRACT_TEMPLATE
    replacements = {
        "___员工姓名___": data.get("员工姓名", "XXX"),
        "___身份证号___": data.get("身份证号", "XXX"),
        "___户籍地址___": data.get("户籍地址", "XXX"),
        "___联系地址___": data.get("联系地址", "XXX"),
        "___手机号___": data.get("手机号", "XXX"),
        "___岗位名称___": data.get("岗位名称", "XXX"),
        "___税前工资___": data.get("税前工资", "XXX"),
        "___工作地点___": data.get("工作地点", "北京市"),
        "___签订年份___": sign_year,
        "___签订月份___": sign_month,
        "___签订日期___": sign_day,
        "___结束年份___": end_year,
        "___结束月份___": end_month,
        "___结束日期___": end_day,
        "___合同年限___": data.get("合同年限", "3"),
        "___试用期月数___": data.get("试用期月数", "3"),
        "___岗位职责描述___": data.get("岗位职责描述", "XXX"),
    }
    
    for placeholder, value in replacements.items():
        content = content.replace(placeholder, value)
    
    # 保存为Word文档
    if not output_path:
        filename = f"劳动合同_{data.get('员工姓名', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = f"/tmp/{filename}"
    
    # 创建Word文档
    doc = Document()
    
    # 设置字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 逐行添加内容
    for line in content.split('\n'):
        if line.strip():
            # 检测是否是标题
            if line.startswith('劳动合同') and len(line) < 10:
                p = doc.add_heading(line, level=1)
            elif line.startswith('合同期限') or line.startswith('工作内容') or line.startswith('工作时间和') or \
                 line.startswith('四、劳动') or line.startswith('五、岗位') or line.startswith('六、劳动') or \
                 line.startswith('七、劳动') or line.startswith('八、保密') or line.startswith('九、合同') or \
                 line.startswith('十、争议') or line.startswith('十一、其他'):
                p = doc.add_heading(line, level=2)
            else:
                p = doc.add_paragraph(line)
        else:
            doc.add_paragraph()
    
    doc.save(output_path)
    return output_path


# 必填字段
REQUIRED_FIELDS = ["员工姓名", "岗位名称", "税前工资"]


def check_missing_fields(data: dict) -> list:
    """检查缺失的必填字段"""
    missing = []
    for field in REQUIRED_FIELDS:
        if field not in data or not data[field] or data[field] == "XXX":
            missing.append(field)
    return missing


def generate_missing_prompt(missing: list) -> str:
    """生成追问消息"""
    prompts = {
        "员工姓名": "员工叫什么名字？",
        "岗位名称": "是什么岗位？（比如：产品经理、前端开发等）",
        "税前工资": "月薪是多少？"
    }
    return "信息还缺一点哦：\n" + "\n".join([f"• {prompts.get(f, f)}" for f in missing])


if __name__ == "__main__":
    # 测试
    test_data = {
        "员工姓名": "张三",
        "身份证号": "110101199001011234",
        "户籍地址": "北京市海淀区xxx",
        "联系地址": "北京市朝阳区xxx",
        "手机号": "13800138000",
        "岗位名称": "人事实习生",
        "税前工资": "20000",
        "工作地点": "北京市",
        "签订日期": "2026-03-25",
        "岗位职责描述": "负责招聘流程管理、协助员工入职办理、整理人事档案",
        "合同年限": "3",
        "试用期月数": "3"
    }
    
    output = generate_labor_contract(test_data)
    print(f"合同已生成: {output}")
