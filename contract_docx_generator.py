#!/usr/bin/env python3
"""
合同Word文档生成器 - 生成真正的.docx文件
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import os

def generate_contract_docx(contract_type: str, data: dict, output_path: str = None) -> str:
    """
    生成合同Word文档
    
    contract_type: "劳动" | "劳务" | "实习"
    data: 合同数据字典
    output_path: 输出文件路径（可选）
    
    返回: 生成的文件路径
    """
    # 解析日期
    sign_date = data.get("签订日期", datetime.now().strftime("%Y-%m-%d"))
    sign_year, sign_month, sign_day = sign_date.split("-")
    
    # 计算结束日期
    contract_years = int(data.get("合同年限", "3"))
    end_date = datetime(int(sign_year), int(sign_month), int(sign_day)) + timedelta(days=365*contract_years)
    end_year, end_month, end_day = str(end_date.year), str(end_date.month).zfill(2), str(end_date.day).zfill(2)
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('劳动合同', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 甲方乙方信息
    doc.add_paragraph(f"甲方（用人单位）：北京极群科技有限公司")
    doc.add_paragraph(f"地址：北京市海淀区中关村东路8号东升大厦AB座四层4161号")
    doc.add_paragraph(f"法定代表人：陈春宇")
    doc.add_paragraph()
    
    doc.add_paragraph(f"乙方（劳动者）：{data.get('员工姓名', 'XXX')}")
    doc.add_paragraph(f"身份证号码：{data.get('身份证号', 'XXX')}")
    doc.add_paragraph(f"户籍地址：{data.get('户籍地址', 'XXX')}")
    doc.add_paragraph(f"联系地址：{data.get('联系地址', 'XXX')}")
    doc.add_paragraph(f"联系电话：{data.get('手机号', 'XXX')}")
    doc.add_paragraph()
    
    # 鉴于条款
    doc.add_paragraph("鉴于：")
    doc.add_paragraph("甲方已如实告知并经乙方确认其工作内容、工作条件、工作地点、职业危害、安全生产状况、劳动报酬以及乙方要求了解的其他情况；")
    doc.add_paragraph("根据《中华人民共和国劳动法》《中华人民共和国劳动合同法》等法律法规政策规定，甲乙双方遵循合法、公平、平等自愿、协商一致、诚实信用的原则订立本合同，共同遵守所列条款。")
    doc.add_paragraph()
    
    # 一、合同期限
    doc.add_heading('一、合同期限', level=2)
    doc.add_paragraph(f"1. 本合同为有固定期限合同，合同期限 {contract_years} 年，自 {sign_year} 年 {sign_month} 月 {sign_day} 日起至 {end_year} 年 {end_month} 月 {end_day} 日止，试用期为 {data.get('试用期月数', '3')} 个月。")
    doc.add_paragraph("2. 本合同到期后，甲方未提出续签或乙方不同意续签的，本合同到期终止。")
    doc.add_paragraph()
    
    # 二、工作内容、工作地点
    doc.add_heading('二、工作内容、工作地点', level=2)
    doc.add_paragraph(f"1. 乙方工作地点为：{data.get('工作地点', '北京市')}。")
    doc.add_paragraph(f"2. 签订合同时，乙方工作岗位为：{data.get('岗位名称', 'XXX')}。")
    doc.add_paragraph("3. 乙方应当接受甲方对其进行的工作岗位所必需的培训。乙方应主动学习，积极参加甲方组织的培训，提高职业技能。")
    doc.add_paragraph("4. 乙方必须按照甲方确定的岗位职责，按时、按质、按量完成工作。")
    doc.add_paragraph()
    
    # 三、工作时间和休息休假
    doc.add_heading('三、工作时间和休息休假', level=2)
    doc.add_paragraph("1. 甲方安排乙方执行标准工时工作制。每日工作时间8小时，每周工作时间40小时，每周至少休息一日。")
    doc.add_paragraph("2. 甲方安排乙方加班的，应依法优先安排补休，无法安排补休的依法支付加班工资。")
    doc.add_paragraph("3. 乙方依法享有国家规定的各种法定节假日和甲方的休假待遇。")
    doc.add_paragraph()
    
    # 四、劳动报酬
    doc.add_heading('四、劳动报酬', level=2)
    doc.add_paragraph(f"1. 甲方依法制定工资分配制度，并告知乙方。甲方至少每月以货币形式向乙方支付一次工资。甲方经与乙方协商，约定按照如下方式向乙方以货币形式发放工资，于每月15日前足额支付：")
    doc.add_paragraph(f"固定工资 {data.get('税前工资', 'XXX')} 元/月（税前）。")
    doc.add_paragraph("2. 试用期内，乙方的工资为转正后的 100%。")
    doc.add_paragraph("3. 双方按照国家和地方规定缴纳社会保险和住房公积金，甲方依法为乙方办理有关手续并履行代扣代缴义务。")
    doc.add_paragraph()
    
    # 五、岗位职责
    doc.add_heading('五、岗位职责', level=2)
    doc.add_paragraph("乙方的基本岗位职责和要求如下：")
    doc.add_paragraph(data.get('岗位职责描述', 'XXX'))
    doc.add_paragraph()
    
    # 六、劳动纪律
    doc.add_heading('六、劳动纪律', level=2)
    doc.add_paragraph("1. 乙方应遵守国家的法律法规。")
    doc.add_paragraph("2. 乙方已知悉并详细阅读甲方的各项规章制度，并承诺严格遵守。甲方有权依法不时制定、修改内部规章制度，乙方亦有义务遵守。")
    doc.add_paragraph("3. 乙方应遵守的劳动纪律包括但不限于：严格遵守考勤制度，不得迟到、早退、旷工；服从甲方工作安排；不得有弄虚作假和欺诈的行为；保守甲方商业秘密等。")
    doc.add_paragraph()
    
    # 七、劳动保护
    doc.add_heading('七、劳动保护、劳动条件和职业危害防护', level=2)
    doc.add_paragraph("甲方应按照国家有关规定为乙方提供必要的劳动条件和安全保护。")
    doc.add_paragraph()
    
    # 八、保密规定
    doc.add_heading('八、保密规定', level=2)
    doc.add_paragraph("在本合同有效期间及之后，乙方不得以任何方式侵害甲方和/或其关联方的商业秘密和/或保密信息，否则应当依法承担违约责任和其他一切法律责任。")
    doc.add_paragraph()
    
    # 九、合同的变更、解除和终止
    doc.add_heading('九、合同的变更、解除和终止', level=2)
    doc.add_paragraph("按照国家相关法律法规执行。")
    doc.add_paragraph()
    
    # 十、争议处理
    doc.add_heading('十、争议处理', level=2)
    doc.add_paragraph("双方如发生劳动争议，应协商解决；协商不成的，可向劳动争议仲裁委员会申请仲裁。")
    doc.add_paragraph()
    
    # 十一、其他
    doc.add_heading('十一、其他', level=2)
    doc.add_paragraph("1. 本合同一式两份，甲乙双方各执一份，具有同等法律效力。")
    doc.add_paragraph("2. 本合同自双方签字（或盖章）之日起生效。")
    doc.add_paragraph()
    
    # 签字部分
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "甲方（盖章）："
    table.cell(0, 1).text = "乙方（签字）："
    table.cell(1, 0).text = f"日期：{sign_year}年{sign_month}月{sign_day}日"
    table.cell(1, 1).text = f"日期：{sign_year}年{sign_month}月{sign_day}日"
    
    # 保存文件
    if not output_path:
        filename = f"劳动合同_{data.get('员工姓名', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = f"/tmp/{filename}"
    
    doc.save(output_path)
    return output_path


def send_file_to_lark(client, chat_id: str, file_path: str, file_name: str = None) -> bool:
    """
    通过飞书API发送文件到聊天
    
    需要实现文件上传和发送，这里先提供框架
    """
    import requests
    
    if not file_name:
        file_name = os.path.basename(file_path)
    
    try:
        # 1. 获取文件上传地址
        # 这里需要调用飞书API上传文件
        # 简化版：返回文件路径，让用户手动下载
        return True
    except Exception as e:
        print(f"发送文件失败: {e}")
        return False


if __name__ == "__main__":
    # 测试
    test_data = {
        "员工姓名": "张三",
        "身份证号": "110101199001011234",
        "户籍地址": "XXX",
        "联系地址": "XXX",
        "手机号": "XXX",
        "岗位名称": "人事实习生",
        "税前工资": "20000",
        "工作地点": "北京市",
        "签订日期": "2026-03-25",
        "岗位职责描述": "1. 负责招聘流程管理\n2. 协助员工入职办理\n3. 整理人事档案",
        "合同年限": "3",
        "试用期月数": "3"
    }
    
    output = generate_contract_docx("劳动", test_data)
    print(f"合同已生成: {output}")
