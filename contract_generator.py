#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
contract_generator.py  ——  合同生成模块 v2.0
北京极群科技有限公司

支持三种合同类型，基于真实 DOCX 模板精确填充，格式原样保留。
对外接口：
  generate_contract(contract_type, fields, output_name=None) -> str  生成文件路径
  extract_fields_via_llm(user_message, llm_client) -> (str, dict)   LLM提取字段
  detect_contract_type(text) -> str                                  识别合同类型
"""

import os
import re
import json
import logging
from copy import deepcopy
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ─── 路径配置 ──────────────────────────────────────────────────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.environ.get("CONTRACT_TEMPLATES_DIR",
                               os.path.join(_HERE, "templates"))
OUTPUT_DIR    = os.environ.get("CONTRACT_OUTPUT_DIR", "/tmp")

TEMPLATES = {
    "labor":   os.path.join(TEMPLATES_DIR, "中国大陆-劳动合同-空模板.docx"),
    "service": os.path.join(TEMPLATES_DIR, "中国大陆-劳务合同-模板-20260109draft.docx"),
    "intern":  os.path.join(TEMPLATES_DIR, "中国大陆-实习合同-模板.docx"),
}

CONTRACT_TYPE_NAMES = {
    "labor":   "劳动合同",
    "service": "劳务合同",
    "intern":  "实习合同",
}


# ─── 通用工具 ──────────────────────────────────────────────────────────────────

def parse_date(date_str):
    if not date_str:
        return None
    s = str(date_str).strip()
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            return datetime.strptime(s, fmt)
        except ValueError:
            pass
    m = re.match(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', s)
    if m:
        return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    return None


def calc_end_date(start, duration, unit):
    if not start or not duration:
        return None
    try:
        n = int(str(duration))
    except ValueError:
        return None
    return start + relativedelta(years=n if unit == '年' else 0,
                                  months=0 if unit == '年' else n)


def fmt_y(d): return str(d.year)
def fmt_m(d): return f'{d.month:02d}'
def fmt_d(d): return f'{d.day:02d}'


def _replace_run_at(para, idx, text):
    if 0 <= idx < len(para.runs):
        para.runs[idx].text = text
        return True
    return False


def _set_para_text(para, new_text):
    """将段落内所有 run 合并到第一个 run，保留字符格式。"""
    runs = para.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ''
    else:
        para.add_run(new_text)


def _delete_para(para):
    p = para._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _clone_pPr(src_para, dst_para):
    src_pPr = src_para._p.find(qn('w:pPr'))
    if src_pPr is None:
        return
    dst_pPr = dst_para._p.find(qn('w:pPr'))
    if dst_pPr is not None:
        dst_para._p.remove(dst_pPr)
    dst_para._p.insert(0, deepcopy(src_pPr))


def _clone_rPr(src_run, dst_run):
    src_rPr = src_run._r.find(qn('w:rPr'))
    if src_rPr is None:
        return
    dst_rPr = dst_run._r.find(qn('w:rPr'))
    if dst_rPr is not None:
        dst_run._r.remove(dst_rPr)
    dst_run._r.insert(0, deepcopy(src_rPr))


def _parse_responsibilities(resp_input):
    """
    解析岗位职责，返回 [{"text": str, "bold": bool}, ...]
    支持 str、list[str]、list[dict]
    """
    if not resp_input:
        return []
    items = []
    if isinstance(resp_input, str):
        lines = resp_input.split('\n')
    elif isinstance(resp_input, list):
        lines = []
        for item in resp_input:
            if isinstance(item, dict):
                items.append(item)
            else:
                lines.append(str(item))
    else:
        lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        first = stripped[0]
        is_bold = not (first in '（(- \t' or first.isdigit())
        items.append({"text": stripped, "bold": is_bold})
    return items


# ─── 劳动合同 ──────────────────────────────────────────────────────────────────

def _fill_labor(doc, fields):
    paras = doc.paragraphs
    sign_date  = parse_date(fields.get('sign_date'))
    start_date = parse_date(fields.get('start_date'))
    end_date   = parse_date(fields.get('end_date'))
    duration   = str(fields.get('duration', ''))
    dur_unit   = fields.get('duration_unit', '年')
    if not end_date and start_date and duration:
        end_date = calc_end_date(start_date, duration, dur_unit)
    probation  = str(fields.get('probation_period', '3'))
    name       = str(fields.get('name', ''))
    id_num     = str(fields.get('id_number', ''))
    contact    = str(fields.get('contact_address', ''))
    household  = str(fields.get('household_address', ''))
    phone      = str(fields.get('phone', ''))
    salary     = str(fields.get('salary', ''))
    job_title  = str(fields.get('job_title', ''))
    work_loc   = str(fields.get('work_location', '北京市'))

    if name:       _replace_run_at(paras[8], 3, name)
    if sign_date:
        _replace_run_at(paras[9], 2, fmt_y(sign_date))
        _replace_run_at(paras[9], 6, fmt_m(sign_date))
        _replace_run_at(paras[9], 10, fmt_d(sign_date))
        _replace_run_at(paras[11], 6, fmt_y(sign_date))
        _replace_run_at(paras[11], 10, fmt_m(sign_date))
        r14 = paras[11].runs[14] if len(paras[11].runs) > 14 else None
        if r14:
            r14.text = fmt_d(sign_date) + (' ' if r14.text.endswith(' ') else '')
    if duration:   _replace_run_at(paras[20], 6, duration)
    if start_date:
        _replace_run_at(paras[20], 10, fmt_y(start_date))
        _replace_run_at(paras[20], 14, fmt_m(start_date))
        _replace_run_at(paras[20], 18, fmt_d(start_date))
    if end_date:
        _replace_run_at(paras[20], 22, fmt_y(end_date))
        _replace_run_at(paras[20], 26, fmt_m(end_date))
        _replace_run_at(paras[20], 30, fmt_d(end_date))
    if probation:  _replace_run_at(paras[20], 34, probation)
    if work_loc:   _replace_run_at(paras[25], 3, work_loc)
    if job_title:  _replace_run_at(paras[26], 3, job_title)
    if salary:     _replace_run_at(paras[40], 6, salary)
    if job_title:  _replace_run_at(paras[113], 2, job_title)

    # 表格：单行双列，右列 para[0..5]
    cell = doc.tables[1].rows[0].cells[1]
    if name:      _set_para_text(cell.paragraphs[0], name)
    if id_num:    _set_para_text(cell.paragraphs[1], id_num)
    if household: _set_para_text(cell.paragraphs[3], household)
    if contact:   _set_para_text(cell.paragraphs[4], contact)
    if phone:     _set_para_text(cell.paragraphs[5], phone)

    # 附件职责
    resp_items = _parse_responsibilities(fields.get('job_responsibilities'))
    if resp_items:
        ref_n_para = paras[115]
        ref_n_run  = paras[115].runs[0]
        ref_b_run  = paras[118].runs[0]
        for p in list(paras[115:]):
            _delete_para(p)
        for item in resp_items:
            np = doc.add_paragraph()
            _clone_pPr(ref_n_para, np)
            nr = np.add_run(item['text'])
            _clone_rPr(ref_b_run if item['bold'] else ref_n_run, nr)
            nr.bold = item['bold']


# ─── 劳务合同 ──────────────────────────────────────────────────────────────────

def _fill_service(doc, fields):
    paras = doc.paragraphs
    sign_date  = parse_date(fields.get('sign_date'))
    start_date = parse_date(fields.get('start_date'))
    end_date   = parse_date(fields.get('end_date'))
    duration   = str(fields.get('duration', ''))
    dur_unit   = fields.get('duration_unit', '月')
    if not end_date and start_date and duration:
        end_date = calc_end_date(start_date, duration, dur_unit)
    name      = str(fields.get('name', ''))
    id_num    = str(fields.get('id_number', ''))
    contact   = str(fields.get('contact_address', ''))
    household = str(fields.get('household_address', ''))
    phone     = str(fields.get('phone', ''))
    salary    = str(fields.get('salary', ''))
    job_title = str(fields.get('job_title', ''))
    work_loc  = str(fields.get('work_location', '北京市'))

    if sign_date:
        _replace_run_at(paras[2], 3, fmt_y(sign_date))
        _replace_run_at(paras[2], 4,
                        f'年  {fmt_m(sign_date)}  月  {fmt_d(sign_date)}  日签订：')
    if duration:   _replace_run_at(paras[9], 1, duration)
    if start_date:
        _replace_run_at(paras[9], 3, fmt_y(start_date))
        _replace_run_at(paras[9], 5, fmt_m(start_date))
        _replace_run_at(paras[9], 7, fmt_d(start_date))
    if end_date:
        _replace_run_at(paras[9],  9, fmt_y(end_date))
        _replace_run_at(paras[9], 11, fmt_m(end_date))
        _replace_run_at(paras[9], 13, fmt_d(end_date))
    if work_loc:   _replace_run_at(paras[12], 2, f'  {work_loc}  ')
    if job_title:  _replace_run_at(paras[13], 1, f'  {job_title}  ')
    if salary:
        _replace_run_at(paras[19], 0, '☑')
        _replace_run_at(paras[19], 2, f'  {salary}  ')

    cell = doc.tables[1].rows[0].cells[1]
    if name:      _set_para_text(cell.paragraphs[0], name)
    if id_num:    _set_para_text(cell.paragraphs[1], id_num)
    if household: _set_para_text(cell.paragraphs[3], household)
    if contact:   _set_para_text(cell.paragraphs[4], contact)
    if phone:     _set_para_text(cell.paragraphs[5], phone)

    if job_title: _replace_run_at(paras[95], 4, f' {job_title}')

    resp_items = _parse_responsibilities(fields.get('job_responsibilities'))
    if resp_items:
        ref_n_para = paras[98]
        ref_n_run  = paras[99].runs[0]
        ref_b_run  = paras[98].runs[0]
        for p in list(paras[98:]):
            _delete_para(p)
        for item in resp_items:
            np = doc.add_paragraph()
            _clone_pPr(ref_n_para, np)
            nr = np.add_run(item['text'])
            _clone_rPr(ref_b_run if item['bold'] else ref_n_run, nr)
            nr.bold = item['bold']


# ─── 实习合同 ──────────────────────────────────────────────────────────────────

def _fill_intern(doc, fields):
    paras = doc.paragraphs
    sign_date  = parse_date(fields.get('sign_date'))
    start_date = parse_date(fields.get('start_date'))
    end_date   = parse_date(fields.get('end_date'))
    duration   = str(fields.get('duration', ''))
    dur_unit   = fields.get('duration_unit', '月')
    if not end_date and start_date and duration:
        end_date = calc_end_date(start_date, duration, dur_unit)
    name        = str(fields.get('name', ''))
    id_num      = str(fields.get('id_number', ''))
    contact     = str(fields.get('contact_address', ''))
    household   = str(fields.get('household_address', ''))
    phone       = str(fields.get('phone', ''))
    salary      = str(fields.get('salary', ''))
    salary_type = str(fields.get('salary_type', 'fixed'))
    job_title   = str(fields.get('job_title', ''))
    work_loc    = str(fields.get('work_location', '北京市'))

    if sign_date:
        _replace_run_at(paras[2], 9, fmt_y(sign_date))
        _replace_run_at(paras[2], 12, f' {fmt_m(sign_date)}')
        _replace_run_at(paras[2], 14, fmt_d(sign_date))
    if duration:   _replace_run_at(paras[9], 2, duration)
    if start_date:
        _replace_run_at(paras[9], 7,  fmt_y(start_date))
        _replace_run_at(paras[9], 11, f' {fmt_m(start_date)}')
        _replace_run_at(paras[9], 15, f' {fmt_d(start_date)}')
    if end_date:
        _replace_run_at(paras[9], 19, fmt_y(end_date))
        _replace_run_at(paras[9], 23, f' {fmt_m(end_date)}')
        _replace_run_at(paras[9], 27, fmt_d(end_date))
    if work_loc:
        _replace_run_at(paras[12], 2, f'  {work_loc}  ')
        for idx in [3, 4, 5, 6, 7]:
            _replace_run_at(paras[12], idx, '')
    if job_title:  _replace_run_at(paras[13], 2, job_title)
    if salary:
        if salary_type == 'daily':
            _replace_run_at(paras[21], 3, salary)
        else:
            _replace_run_at(paras[19], 2, f'  {salary}  ')

    cell = doc.tables[1].rows[0].cells[1]
    if name:      _set_para_text(cell.paragraphs[0], name)
    if id_num:    _set_para_text(cell.paragraphs[1], id_num)
    if household: _set_para_text(cell.paragraphs[3], household)
    if contact:   _set_para_text(cell.paragraphs[4], contact)
    if phone:     _set_para_text(cell.paragraphs[5], phone)

    if job_title:
        _replace_run_at(paras[95], 4, f' {job_title}')
        for idx in [5, 6, 7]:
            _replace_run_at(paras[95], idx, '')

    resp_items = _parse_responsibilities(fields.get('job_responsibilities'))
    if resp_items:
        ref_n_para = paras[97]
        ref_n_run  = paras[97].runs[0]
        for p in list(paras[97:]):
            _delete_para(p)
        for item in resp_items:
            np = doc.add_paragraph()
            _clone_pPr(ref_n_para, np)
            nr = np.add_run(item['text'])
            _clone_rPr(ref_n_run, nr)
            nr.bold = item['bold']


# ─── 公开接口 ──────────────────────────────────────────────────────────────────

def generate_contract(contract_type: str, fields: dict,
                      output_name: str = None) -> str:
    """
    生成合同 Word 文档，返回文件路径。

    contract_type : "labor" | "service" | "intern"
    fields        : 字段字典（见模块文档）
    output_name   : 文件名前缀（默认取 fields["name"]）
    """
    if contract_type not in TEMPLATES:
        raise ValueError(f"未知合同类型: {contract_type}")
    tpl = TEMPLATES[contract_type]
    if not os.path.exists(tpl):
        raise FileNotFoundError(f"模板不存在: {tpl}\n请确认 templates/ 目录已部署")

    doc = Document(tpl)
    if contract_type == "labor":
        _fill_labor(doc, fields)
    elif contract_type == "service":
        _fill_service(doc, fields)
    elif contract_type == "intern":
        _fill_intern(doc, fields)

    # 安全保险：确保乙方签字行保持留白（不含人名）
    _name = fields.get("name", "")
    if _name:
        _sign_kws = ["签字", "签名", "签署", "乙方（劳动者）", "乙方（劳务人员）",
                     "乙方（实习生）", "乙方签", "乙  方", "乙方："]
        # 正文段落 + 所有表格单元格段落都要检查
        all_paras = list(doc.paragraphs)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)
        for para in all_paras:
            if any(kw in para.text for kw in _sign_kws) and _name in para.text:
                new_t = para.text.replace(_name, "")
                _set_para_text(para, new_t)
                logger.warning(f"已清除签字行中的姓名: {para.text[:40]}")

    name = output_name or fields.get('name', '未命名')
    filename = f"{name}-{CONTRACT_TYPE_NAMES[contract_type]}.docx"
    out_path = os.path.join(OUTPUT_DIR, filename)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(out_path)
    logger.info(f"合同已生成: {out_path}")
    return out_path


def detect_contract_type(text: str) -> str:
    """从消息文本中识别合同类型，返回 'labor' / 'service' / 'intern'"""
    if re.search(r'实习合同|实习生合同', text):
        return 'intern'
    if re.search(r'劳务合同|劳务', text):
        return 'service'
    return 'labor'   # 默认：劳动合同


# LLM提取的系统提示词（供 app.py 调用）
CONTRACT_EXTRACT_SYSTEM_PROMPT = """你是一个合同信息提取助手。
从用户消息中提取合同相关字段，以 JSON 格式返回。
只返回纯 JSON，不加解释。

可提取的字段（均为字符串）：
  name               姓名
  id_number          身份证号码（18位）
  contact_address    联系地址（现住址）
  household_address  户籍地址
  phone              联系电话
  sign_date          合同签订日期 YYYY-MM-DD
  start_date         入职/合同开始日期 YYYY-MM-DD
  end_date           合同结束日期 YYYY-MM-DD（可选，与 duration 二选一）
  duration           合同时长数字（如 "3"）
  duration_unit      时长单位 "年" 或 "月"（劳动合同一般填"年"，实习/劳务填"月"）
  probation_period   试用期月数（仅劳动合同，如 "3"）
  salary             薪资数字字符串（如 "20000"）
  salary_type        "fixed"=月薪 / "daily"=日薪（实习合同）
  job_title          岗位名称
  work_location      工作地点（默认"北京市"）
  job_responsibilities  岗位职责（字符串列表，段落标题加粗，正文条目以（1）等开头）

只包含用户消息中明确提到的字段，未提及的字段不要输出。"""


def _extract_fields_regex(text: str) -> dict:
    """
    纯正则兜底提取合同字段（LLM 不可用时使用）。
    支持结构化多行格式（字段名：值）和自然语言混合格式。
    """
    fields = {}

    def _norm_date(s: str) -> str:
        """把各种日期格式统一为 YYYY-MM-DD"""
        s = re.sub(r'[年月]', '-', s).rstrip('日').replace('/', '-').strip()
        parts = s.split('-')
        if len(parts) == 3:
            return f"{parts[0]}-{int(parts[1]):02d}-{int(parts[2]):02d}"
        return s

    # ── 姓名 ──────────────────────────────────────────────────────────────
    # 优先：姓名：张三 / 姓名张三 / 姓名是张三
    m = re.search(r'姓名[是为：:\s]*([^\s，,。！?、\d]{2,5})', text)
    if m:
        fields['name'] = m.group(1).strip()
    # 兜底：X的合同（如"出一份张三的合同"）
    if not fields.get('name'):
        m = re.search(r'([^\s，,、出份一]{2,4})的(?:劳动|劳务|实习)?合同', text)
        if m:
            fields['name'] = m.group(1).strip()

    # ── 岗位 ──────────────────────────────────────────────────────────────
    m = re.search(r'(?:职位|岗位|职务|职称)[是为：:\s]*([^\s，,。！?\d、]{2,15})', text)
    if m:
        fields['job_title'] = m.group(1).strip()
    # 带后缀的岗位（如"产品经理"、"运营专员"）
    if not fields.get('job_title'):
        m = re.search(r'[，,]\s*([^\s，,、\d]{2,8}(?:部|岗|师|员|助理|专员|经理|总监|工程师|实习生|运营|产品|设计|开发))\s*[，,]', text)
        if m:
            fields['job_title'] = m.group(1).strip()
    # 兜底：逗号之间的 2-4 字中文短词（排除已知非岗位词）
    # 用零宽断言避免逗号被消耗，相邻词都能匹配
    if not fields.get('job_title'):
        _non_job = {'实习合同', '劳动合同', '劳务合同', '一天', '每天', '月薪', '日薪'}
        for m in re.finditer(r'(?<=[，,])\s*([^\s，,。！?\d、]{2,5})\s*(?=[，,])', text):
            cand = m.group(1).strip()
            if not any(kw in cand for kw in _non_job):
                fields['job_title'] = cand
                break

    # ── 签订日期 ──────────────────────────────────────────────────────────
    m = re.search(r'签订日期[是为：:\s]*(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}|\d{4}-\d{2}-\d{2})', text)
    if m:
        fields['sign_date'] = _norm_date(m.group(1))

    # ── 合同开始日期 ──────────────────────────────────────────────────────
    m = re.search(r'(?:合同)?开始日期[是为：:\s]*(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}|\d{4}-\d{2}-\d{2})', text)
    if not m:
        m = re.search(r'入职(?:日期|时间)?[是为：:\s]*(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}|\d{4}-\d{2}-\d{2})', text)
    if m:
        fields['start_date'] = _norm_date(m.group(1))

    # ── 薪资 ──────────────────────────────────────────────────────────────
    # 月薪
    m = re.search(r'(?:薪资|工资|月薪|薪酬|底薪)[是为：:\s]*([0-9,，.]+)', text)
    if m:
        fields['salary'] = re.sub(r'[,，]', '', m.group(1))
        fields['salary_type'] = 'fixed'
    # 日薪：200一天 / 200/天 / 200元/天 / 日薪200
    if not fields.get('salary'):
        m = re.search(r'([0-9]+)\s*(?:元\s*)?[/每一]天|日薪[是为：:\s]*([0-9]+)', text)
        if m:
            fields['salary'] = m.group(1) or m.group(2)
            fields['salary_type'] = 'daily'

    # ── 身份证号（18位）─────────────────────────────────────────────────
    m = re.search(r'(?:身份证[号码]?[是为：:\s]*)?([0-9]{17}[0-9Xx])', text)
    if m:
        fields['id_number'] = m.group(1).upper()

    # ── 电话 ──────────────────────────────────────────────────────────────
    m = re.search(r'(?:手机|电话|联系电话)[号码]?[是为：:\s]*([1][3-9][0-9]{9})', text)
    if not m:
        m = re.search(r'\b([1][3-9][0-9]{9})\b', text)
    if m:
        fields['phone'] = m.group(1)

    # ── 户籍地址 ──────────────────────────────────────────────────────────
    m = re.search(r'户籍(?:地址)?[是为：:\s]*([^\n，,。！?]{5,50})', text)
    if m:
        fields['household_address'] = m.group(1).strip()

    # ── 联系地址 ──────────────────────────────────────────────────────────
    m = re.search(r'(?:联系|现)(?:居住)?地址[是为：:\s]*([^\n，,。！?]{5,50})', text)
    if m:
        fields['contact_address'] = m.group(1).strip()

    # ── 工作地点 ──────────────────────────────────────────────────────────
    m = re.search(r'(?:工作地点|工作地|城市)[是为：:\s]*([^\s，,。！?、]{2,10})', text)
    if m:
        fields['work_location'] = m.group(1).strip()

    return fields


def extract_fields_via_llm(user_message: str, llm_client) -> tuple:
    """
    调用 LLM 从用户消息中提取合同字段。

    参数:
        user_message: 用户消息
        llm_client  : 有 _call_api(messages) 方法的 LLM 客户端

    返回:
        (contract_type: str, fields: dict)
    """
    contract_type = detect_contract_type(user_message)
    current_year  = datetime.now().year
    messages = [
        {"role": "system", "content": (
            CONTRACT_EXTRACT_SYSTEM_PROMPT +
            f"\n\n当前年份是 {current_year} 年。用户未明确说明年份时，默认为 {current_year} 年。"
        )},
        {"role": "user", "content": user_message}
    ]
    try:
        resp = llm_client._call_api(messages, tools=None, temperature=0)
        if "error" in resp:
            logger.error(f"LLM extract error: {resp['error']}")
            # LLM 失败时用 regex 兜底
            return contract_type, _extract_fields_regex(user_message)
        content = resp["choices"][0]["message"].get("content", "")
        # 兼容带 markdown 代码块的输出
        content = re.sub(r'^```(?:json)?\s*', '', content.strip())
        content = re.sub(r'\s*```$', '', content)
        fields = json.loads(content)
        # 合并 regex 兜底（LLM 漏掉的字段用 regex 补充）
        regex_fields = _extract_fields_regex(user_message)
        for k, v in regex_fields.items():
            if k not in fields or not fields[k]:
                fields[k] = v
        return contract_type, fields
    except Exception as e:
        logger.error(f"extract_fields_via_llm failed: {e}")
        # 异常时也用 regex 兜底
        return contract_type, _extract_fields_regex(user_message)
